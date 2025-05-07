import os
import uuid
import logging
from typing import Dict, List, Any, Optional
import aiofiles
from fastapi import UploadFile, HTTPException, BackgroundTasks
from datetime import datetime, timedelta
import asyncio
import json
import time
import backoff
from tenacity import retry, stop_after_attempt, wait_exponential, retry_if_exception_type, before_sleep_log
from storage3.exceptions import StorageApiError

from app.services.database_service import DatabaseService
from app.services.storage_service import StorageService, get_storage_service
from app.services.embedding_service import EmbeddingService, get_embedding_service
from app.services.vector_store_service import (
    VectorStoreService,
    get_vector_store_service,
)
from app.config.settings import settings

# Configure logging
logger = logging.getLogger(__name__)


class DocumentService:
    """Service for document-related operations including uploads, processing and indexing."""

    def __init__(self):
        """Initialize the document service with required services."""
        self.db_service = DatabaseService()
        self.storage_service = get_storage_service()
        self.embedding_service = get_embedding_service()
        self.vector_store_service = get_vector_store_service()
        self.processing_tasks = {}  # Track processing tasks by document_id
        logger.info("Document service initialized")

    @retry(
        stop=stop_after_attempt(5),
        wait=wait_exponential(multiplier=1, min=2, max=30),
        retry=(
            retry_if_exception_type(HTTPException) |
            retry_if_exception_type(StorageApiError) |
            retry_if_exception_type(Exception)
        ),
        reraise=True,
        before_sleep=before_sleep_log(logger, logging.WARNING),
    )
    async def _retry_storage_operation(self, operation, **kwargs):
        """Retry a storage operation with exponential backoff."""
        try:
            # Call the operation and return its result
            result = await operation(**kwargs)
            return result
        except HTTPException as http_ex:
            logger.warning(f"HTTP error in storage operation: {http_ex}")
            raise
        except StorageApiError as storage_ex:
            logger.warning(f"Storage API error: {storage_ex}")
            raise
        except Exception as e:
            logger.warning(f"Unexpected error in storage operation: {e}")
            raise

    async def process_document_upload(
        self,
        file: UploadFile,
        project_id: str,
        user_id: str,
        name: str = None,
        description: str = None
    ) -> dict:
        """Process a document upload and create records."""
        try:
            # 1. Process file and upload to storage
            try:
                # Use provided name or original filename
                file_name = name or file.filename
                file_extension = file_name.split(".")[-1].lower() if "." in file_name else ""
                
                # Generate a unique ID for the file to prevent name collisions in storage
                unique_id = uuid.uuid4().hex
                storage_path = f"{project_id}/{unique_id}-{file_name}"
                storage_bucket = "documents"
                
                logger.info(f"Uploading file to storage: {storage_path}")
                
                # Read file content
                content = await file.read()
                
                # Content size for logging and database
                content_size = len(content)
                logger.info(f"File size: {content_size / 1024:.2f}KB")
                
                # Upload to storage and capture the result
                upload_result = await self._retry_storage_operation(
                    self.storage_service.upload_document,
                    file_content=content,
                    storage_path=storage_path,
                    storage_bucket=storage_bucket,
                    content_type=file.content_type
                )
                
                # Extract the actual storage path from the upload result
                actual_storage_path = upload_result.get('key', storage_path)
                logger.info(f"File uploaded to storage: {actual_storage_path}")
                
                # Verify file exists in storage using the actual path
                file_exists = await self._check_file_exists(storage_bucket, actual_storage_path)
                if not file_exists:
                    logger.warning(f"File upload verification failed. Retrying verification...")
                    # Wait and retry once
                    await asyncio.sleep(2)
                    file_exists = await self._check_file_exists(storage_bucket, actual_storage_path)
                    
                    if not file_exists:
                        logger.error(f"File upload verification failed after retry: {actual_storage_path}")
                        raise ValueError("File upload failed. File not found in storage after upload.")
                
                logger.info(f"File upload verified: {actual_storage_path}")
                
            except Exception as upload_error:
                logger.error(f"Error uploading file to storage: {str(upload_error)}", exc_info=True)
                # Clean up any partially uploaded file
                try:
                    await self.storage_service.delete_document(path=actual_storage_path, bucket=storage_bucket)
                    logger.info(f"Cleaned up partial upload: {actual_storage_path}")
                except Exception as cleanup_error:
                    logger.error(f"Error cleaning up partial upload: {str(cleanup_error)}")
                
                # Re-raise the original error
                raise ValueError(f"File upload failed: {str(upload_error)}")
            
            # 2. Create document record in database using the actual storage path
            try:
                document = await self.db_service.create_document(
                    name=file_name,
                    description=description,
                    project_id=project_id,
                    user_id=user_id,
                    storage_path=actual_storage_path,  # Use the actual path returned from storage service
                    storage_bucket=storage_bucket,
                    file_type=file_extension,
                    file_size=content_size,
                    metadata={
                        "original_filename": file_name,
                        "upload_timestamp": datetime.utcnow().isoformat(),
                        "content_type": file.content_type
                    }
                )
                
                logger.info(f"Created document record in database: {document['id']}")
                
            except Exception as db_error:
                logger.error(f"Error creating document record: {str(db_error)}", exc_info=True)
                
                # Clean up storage if database creation fails
                try:
                    await self.storage_service.delete_document(path=actual_storage_path, bucket=storage_bucket)
                    logger.info(f"Cleaned up storage after database error: {actual_storage_path}")
                except Exception as cleanup_error:
                    logger.error(f"Error cleaning up storage: {str(cleanup_error)}")
                
                # Re-raise error
                raise ValueError(f"Failed to create document record: {str(db_error)}")
            
            # 3. Queue document for processing
            try:
                # Launch processing in background task
                background_tasks = BackgroundTasks()
                background_tasks.add_task(self.process_document, document["id"])
                
                logger.info(f"Queued document {document['id']} for background processing")
                
                # Return document record with storage details
                return {
                    **document,
                    "storage": {
                        "path": actual_storage_path,
                        "bucket": storage_bucket
                    }
                }
                
            except Exception as queue_error:
                logger.error(f"Error queuing document for processing: {str(queue_error)}", exc_info=True)
                
                # Don't clean up here - document record exists and can be retried
                raise ValueError(f"Failed to queue document for processing: {str(queue_error)}")

        except Exception as e:
            logger.error(f"Document upload processing failed: {str(e)}", exc_info=True)
            raise

    async def process_document_with_timeout(self, document_id: str) -> None:
        """Process a document with timeout protection."""
        timeout = settings.DOCUMENT_PROCESSING_TIMEOUT
        try:
            # Set a timeout to prevent processing from hanging indefinitely
            await asyncio.wait_for(
                self.process_document(document_id), 
                timeout=timeout
            )
            logger.info(f"Document {document_id} processed successfully within {timeout}s timeout")
        except asyncio.TimeoutError:
            error_msg = f"Document processing timed out after {timeout} seconds"
            logger.error(f"[DocID: {document_id}] {error_msg}")
            await self.db_service.update_document(
                document_id, {"status": "failed", "processing_error": error_msg}
            )
        except Exception as e:
            error_msg = f"Document processing failed: {str(e)}"
            logger.error(f"[DocID: {document_id}] {error_msg}")
            await self.db_service.update_document(
                document_id, {"status": "failed", "processing_error": error_msg}
            )
        finally:
            # Clean up the task reference
            if document_id in self.processing_tasks:
                del self.processing_tasks[document_id]

    async def queue_document_processing(self, document_id: str) -> None:
        """Queue a document for processing."""
        logger.info(f"Queuing document for processing: {document_id}")
        
        # Check if document already has a processing task
        if document_id in self.processing_tasks:
            if not self.processing_tasks[document_id].done():
                logger.warning(f"Document {document_id} is already being processed")
                return
            else:
                # Task is done, clean it up
                del self.processing_tasks[document_id]
                
        # Create and store the processing task
        self.processing_tasks[document_id] = asyncio.create_task(
            self.process_document_with_timeout(document_id)
        )
        logger.info(f"Document {document_id} queued for processing")

    async def process_document(self, document_id: str) -> None:
        """
        Process a document for semantic search indexing.

        This method:
        1. Retrieves the document from database
        2. Downloads the file from storage
        3. Extracts text content based on file type
        4. Chunks the text
        5. Creates embeddings
        6. Stores embeddings in vector database
        7. Updates document status
        """
        logger.info(f"[DocID: {document_id}] Starting document processing.")
        processing_start_time = datetime.utcnow()

        # Update document status to processing
        try:
            await self.db_service.update_document(
                document_id, {"status": "processing", "processing_error": None}
            )
            logger.info(f"[DocID: {document_id}] Status set to 'processing'.")
        except Exception as e:
             logger.error(f"[DocID: {document_id}] Failed to update status to processing: {e}")
             return 

        # Store variables that might be needed for cleanup
        storage_path = None
        storage_bucket = None

        try:
            # 1. Retrieve document from database
            document = await self.db_service.get_document(document_id)
            if not document:
                error_msg = "Document not found in database."
                logger.error(f"[DocID: {document_id}] {error_msg}")
                await self.db_service.update_document(
                    document_id, {"status": "failed", "processing_error": error_msg}
                )
                return

            project_id = document.get("project_id")
            storage_path = document.get("storage_path")
            storage_bucket = document.get("storage_bucket")
            file_type = document.get("file_type")

            # Set namespace for vector storage - using a consistent pattern for namespaces
            vector_namespace = f"proj_{project_id}"
            logger.info(f"[DocID: {document_id}] Using vector namespace: {vector_namespace}")

            # 2. Verify file exists before attempting to download
            try:
                logger.info(f"[DocID: {document_id}] Verifying file exists before download: {storage_path}")
                file_exists = await self._check_file_exists(storage_bucket, storage_path)
                if not file_exists:
                    # Wait and retry once more with a longer timeout in case file is still being processed
                    logger.warning(f"[DocID: {document_id}] File not found on first check, waiting 10 seconds and retrying...")
                    await asyncio.sleep(10)
                    file_exists = await self._check_file_exists(storage_bucket, storage_path)
                    
                    if not file_exists:
                        error_msg = "File not found in storage after retrying."
                        logger.error(f"[DocID: {document_id}] {error_msg}")
                        await self.db_service.update_document(document_id, {
                            "status": "failed", 
                            "processing_error": error_msg
                        })
                        return
                
                logger.info(f"[DocID: {document_id}] File existence verified, proceeding with download.")
            except Exception as e:
                error_msg = f"File existence check failed: {str(e)}"
                logger.error(f"[DocID: {document_id}] {error_msg}")
                await self.db_service.update_document(document_id, {
                    "status": "failed", 
                    "processing_error": error_msg
                })
                return
            
            # 3. Download file from storage
            logger.info(f"[DocID: {document_id}] Downloading document from storage: {storage_path}")
            try:
                file_content = await self._retry_storage_operation(
                    self.storage_service.get_document,
                    path=storage_path,
                    bucket=storage_bucket
                )
                
                if not file_content:
                    error_msg = "Downloaded file was empty or null. The file may be corrupted."
                    logger.error(f"[DocID: {document_id}] {error_msg}")
                    await self.db_service.update_document(document_id, {
                        "status": "failed", 
                        "processing_error": error_msg
                    })
                    return
                    
                # Update document with file size
                file_size = len(file_content)
                logger.info(f"[DocID: {document_id}] File downloaded successfully ({file_size} bytes).")
                await self.db_service.update_document(document_id, {
                    "file_size": file_size,
                    "processing_progress": 25  # 25% progress - download complete
                })
                
            except Exception as e:
                error_type = type(e).__name__
                error_msg = str(e)
                
                if "404" in error_msg or "not found" in error_msg.lower() or "not_found" in error_msg:
                    error_detail = f"File not found in storage: {storage_bucket}/{storage_path}. The file may have been deleted or was not properly uploaded."
                elif "access denied" in error_msg.lower() or "forbidden" in error_msg.lower() or "unauthorized" in error_msg.lower():
                    error_detail = f"Access denied to file in storage: {storage_bucket}/{storage_path}. Please check storage permissions."
                else:
                    error_detail = f"Failed to download document from storage: {error_type} - {error_msg}"
                    
                logger.error(f"[DocID: {document_id}] {error_detail}", exc_info=True)
                await self.db_service.update_document(document_id, {
                    "status": "failed", 
                    "processing_error": error_detail[:1000]  # Limit error length
                })
                return

            # 4. Extract text based on file type with retry
            logger.info(f"[DocID: {document_id}] Extracting text for file type: {file_type}")
            text_content = await self._retry_extraction(
                self._extract_text_from_file,
                file_content=file_content, 
                file_type=file_type, 
                document_id=document_id
            )
            
            if not text_content:
                error_msg = f"Failed to extract text from {file_type} file or file is empty."
                logger.error(f"[DocID: {document_id}] {error_msg}")
                await self.db_service.update_document(document_id, {
                    "status": "failed", 
                    "processing_error": error_msg
                })
                return
                
            logger.info(f"[DocID: {document_id}] Extracted {len(text_content)} characters of text.")

            # 5. Chunk the text using settings for size and overlap
            logger.info(f"[DocID: {document_id}] Chunking text content.")
            chunks = self._chunk_text(
                text_content, 
                chunk_size=settings.CHUNK_SIZE, 
                overlap=settings.CHUNK_OVERLAP
            )
            
            if not chunks:
                error_msg = "Failed to chunk text content or content is empty."
                logger.error(f"[DocID: {document_id}] {error_msg}")
                await self.db_service.update_document(document_id, {
                    "status": "failed", 
                    "processing_error": error_msg
                })
                return
                
            logger.info(f"[DocID: {document_id}] Split text into {len(chunks)} chunks.")

            # 6. Generate embeddings with retry
            logger.info(f"[DocID: {document_id}] Generating embeddings for {len(chunks)} chunks.")
            embeddings = await self._retry_embedding_generation(
                self.embedding_service.generate_embeddings,
                chunks
            )
            
            if not embeddings or len(embeddings) != len(chunks):
                error_msg = f"Failed to generate embeddings. Expected {len(chunks)}, got {len(embeddings) if embeddings else 0}."
                logger.error(f"[DocID: {document_id}] {error_msg}")
                await self.db_service.update_document(document_id, {
                    "status": "failed", 
                    "processing_error": error_msg
                })
                return
                
            logger.info(f"[DocID: {document_id}] Generated {len(embeddings)} embeddings.")

            # 7. Store embeddings in vector database with retry
            logger.info(f"[DocID: {document_id}] Storing {len(embeddings)} vectors in Pinecone.")
            metadata_base = {
                "document_id": document_id,
                "project_id": project_id,
                "file_name": document.get("name", ""),
                "processed_at": datetime.utcnow().isoformat(),
            }
            
            upsert_result = await self._retry_vector_operation(
                self.vector_store_service.upsert_embeddings_with_metadata,
                embeddings=embeddings,
                texts=chunks,
                metadata_base=metadata_base,
                namespace=vector_namespace
            )
            
            if not upsert_result:
                error_msg = "Failed to store vectors in Pinecone."
                logger.error(f"[DocID: {document_id}] {error_msg}")
                await self.db_service.update_document(document_id, {
                    "status": "failed", 
                    "processing_error": error_msg
                })
                return
                
            logger.info(f"[DocID: {document_id}] Stored vectors in Pinecone: {upsert_result}")

            # 8. Update document status to completed
            await self.db_service.update_document(
                document_id,
                {
                    "status": "completed",
                    "chunk_count": len(chunks),
                    "processing_error": None,
                    "pinecone_namespace": vector_namespace
                },
            )
            
            processing_time = (datetime.utcnow() - processing_start_time).total_seconds()
            logger.info(
                f"[DocID: {document_id}] Document processing completed successfully in {processing_time:.2f} seconds."
            )

        except Exception as e:
            error_msg = f"Processing failed: {str(e)}"
            logger.error(f"[DocID: {document_id}] {error_msg}", exc_info=True)
            
            # Attempt to clean up any orphaned storage files if we know the path
            if storage_path and storage_bucket:
                try:
                    await self.storage_service.delete_document(path=storage_path, bucket=storage_bucket)
                    logger.info(f"[DocID: {document_id}] Cleaned up storage for failed document: {storage_path}")
                except Exception as cleanup_error:
                    logger.error(f"[DocID: {document_id}] Failed to clean up storage: {cleanup_error}")
            
            # Update document with detailed error info
            await self.db_service.update_document(
                document_id, {
                    "status": "failed", 
                    "processing_error": error_msg[:1000]  # Limit error message length
                }
            )
    
    @retry(
        stop=stop_after_attempt(3),
        wait=wait_exponential(multiplier=1, min=1, max=10),
        retry=retry_if_exception_type(Exception)
    )
    async def _retry_extraction(self, operation, **kwargs):
        """Retry text extraction operations with exponential backoff."""
        try:
            return await operation(**kwargs)
        except Exception as e:
            logger.warning(f"Text extraction failed: {str(e)}. Retrying...")
            raise
    
    @retry(
        stop=stop_after_attempt(3),
        wait=wait_exponential(multiplier=1, min=2, max=15),
        retry=retry_if_exception_type(Exception)
    )
    async def _retry_embedding_generation(self, operation, chunks):
        """Retry embedding generation with exponential backoff."""
        try:
            return await operation(chunks)
        except Exception as e:
            logger.warning(f"Embedding generation failed: {str(e)}. Retrying...")
            raise
    
    @retry(
        stop=stop_after_attempt(3),
        wait=wait_exponential(multiplier=1, min=1, max=10),
        retry=retry_if_exception_type(Exception)
    )
    async def _retry_vector_operation(self, operation, **kwargs):
        """Retry vector store operations with exponential backoff."""
        try:
            return await operation(**kwargs)
        except Exception as e:
            logger.warning(f"Vector operation failed: {str(e)}. Retrying...")
            raise

    async def _extract_text_from_file(
        self, file_content: bytes, file_type: str, document_id: str
    ) -> str:
        """Extract text content from a file based on its type."""
        try:
            logger.info(
                f"[DocID: {document_id}] Extracting text from {file_type} file"
            )

            if file_type == "txt" or file_type == "md":
                # For text files, just decode the content
                return file_content.decode("utf-8", errors="replace")

            elif file_type == "pdf":
                # Use PyPDF2 to extract text from PDF
                import io
                import PyPDF2
                
                text = ""
                try:
                    pdf_file = io.BytesIO(file_content)
                    pdf_reader = PyPDF2.PdfReader(pdf_file)
                    
                    # Track progress through pages
                    total_pages = len(pdf_reader.pages)
                    logger.info(f"[DocID: {document_id}] PDF has {total_pages} pages")
                    
                    # Process each page with error handling for individual pages
                    for page_num in range(total_pages):
                        try:
                            page = pdf_reader.pages[page_num]
                            page_text = page.extract_text() or ""
                            if page_text:
                                text += page_text + "\n\n"
                            else:
                                logger.warning(f"[DocID: {document_id}] Empty text on PDF page {page_num+1}/{total_pages}")
                        except Exception as page_error:
                            logger.warning(f"[DocID: {document_id}] Error on PDF page {page_num+1}: {page_error}")
                            # Continue with next page instead of failing entire document
                    
                    # Check if we got any text at all
                    if not text.strip():
                        logger.warning(f"[DocID: {document_id}] No text extracted from PDF")
                        # Try alternative extraction method as fallback
                        try:
                            from pdfminer.high_level import extract_text as pdfminer_extract
                            logger.info(f"[DocID: {document_id}] Trying alternative PDF extraction with pdfminer")
                            pdf_file.seek(0)  # Reset file pointer
                            text = pdfminer_extract(pdf_file)
                            
                            if text.strip():
                                logger.info(f"[DocID: {document_id}] Alternative PDF extraction successful")
                            else:
                                logger.warning(f"[DocID: {document_id}] Alternative PDF extraction also yielded no text")
                                return f"Could not extract readable text from this PDF. The document may be scan-based or contain only images. Document ID: {document_id}"
                        except Exception as fallback_error:
                            logger.error(f"[DocID: {document_id}] Fallback PDF extraction failed: {fallback_error}")
                            return f"PDF text extraction failed. The PDF may be scan-based, password-protected, or corrupted. Document ID: {document_id}"
                    
                    logger.info(f"[DocID: {document_id}] Extracted {len(text)} characters from PDF")
                    return text
                    
                except Exception as pdf_error:
                    logger.error(f"[DocID: {document_id}] PDF extraction error: {pdf_error}", exc_info=True)
                    return f"Error extracting PDF: {str(pdf_error)}. Document ID: {document_id}"

            elif file_type == "docx":
                # Use python-docx to extract text from DOCX
                import io
                from docx import Document
                
                try:
                    docx_file = io.BytesIO(file_content)
                    doc = Document(docx_file)
                    
                    # Extract text with paragraph structure preserved
                    paragraphs = []
                    for para in doc.paragraphs:
                        if para.text.strip():  # Only add non-empty paragraphs
                            paragraphs.append(para.text)
                    
                    # Also extract text from tables
                    for table in doc.tables:
                        for row in table.rows:
                            row_text = []
                            for cell in row.cells:
                                if cell.text.strip():
                                    row_text.append(cell.text.strip())
                            if row_text:
                                paragraphs.append(" | ".join(row_text))
                    
                    text = "\n\n".join(paragraphs)
                    
                    if not text.strip():
                        logger.warning(f"[DocID: {document_id}] No text extracted from DOCX")
                        return f"No readable text found in DOCX document. Document ID: {document_id}"
                    
                    logger.info(f"[DocID: {document_id}] Extracted {len(text)} characters from DOCX with {len(paragraphs)} paragraphs")
                    return text
                    
                except Exception as docx_error:
                    logger.error(f"[DocID: {document_id}] DOCX extraction error: {docx_error}", exc_info=True)
                    return f"Error extracting DOCX: {str(docx_error)}. Document ID: {document_id}"

            elif file_type == "csv":
                # For CSV, convert to text with a simple table representation
                try:
                    import pandas as pd
                    import io
                    
                    try:
                        # Try to load as CSV using pandas for better formatting
                        df = pd.read_csv(io.BytesIO(file_content))
                        text = df.to_string(index=False)
                        logger.info(f"[DocID: {document_id}] Parsed CSV with {len(df)} rows and {len(df.columns)} columns")
                    except:
                        # Fallback to simple decode
                        text = file_content.decode("utf-8", errors="replace")
                        
                    return text
                except Exception as csv_error:
                    logger.error(f"[DocID: {document_id}] CSV parsing error: {csv_error}")
                    return file_content.decode("utf-8", errors="replace")

            elif file_type == "json":
                # For JSON, pretty-print to text
                try:
                    json_data = json.loads(
                        file_content.decode("utf-8", errors="replace")
                    )
                    return json.dumps(json_data, indent=2)
                except json.JSONDecodeError:
                    logger.warning(f"[DocID: {document_id}] Invalid JSON content")
                    return file_content.decode("utf-8", errors="replace")

            else:
                raise ValueError(
                    f"Unsupported file type for text extraction: {file_type}"
                )

        except Exception as e:
            logger.error(f"[DocID: {document_id}] Error extracting text from {file_type} file: {str(e)}", exc_info=True)
            raise

    def _chunk_text(
        self, text: str, chunk_size: int = 1000, overlap: int = 200
    ) -> List[str]:
        """Split text into overlapping chunks of specified size."""
        from app.services.embedding_service import chunk_text
        
        # Use the improved chunking implementation from embedding_service
        return chunk_text(text, chunk_size, overlap)

    async def delete_document_embeddings(self, document_id: str) -> None:
        """Delete all embeddings for a document."""
        try:
            # First get the document to determine its project
            document = await self.db_service.get_document(document_id)
            if not document:
                logger.warning(f"Document not found for deletion: {document_id}")
                return
                
            project_id = document.get("project_id")
            namespace = f"proj_{project_id}"
            
            logger.info(f"Deleting embeddings for document: {document_id}, namespace: {namespace}")
            
            # Delete vectors by document ID within the project namespace
            result = await self.vector_store_service.delete_document_vectors(
                document_id=document_id,
                namespace=namespace
            )
            
            logger.info(f"Successfully deleted embeddings for document {document_id}: {result}")
            
        except Exception as e:
            logger.error(f"Error deleting embeddings for document {document_id}: {str(e)}")
            raise

    async def delete_document_from_storage(self, bucket: str, path: str) -> None:
        """Delete a document file from storage."""
        try:
            logger.info(f"Deleting document from storage: {bucket}/{path}")
            await self.storage_service.delete_document(path=path, bucket=bucket)
            logger.info(f"Successfully deleted document from storage: {bucket}/{path}")
        except Exception as e:
            logger.error(
                f"Error deleting document from storage {bucket}/{path}: {str(e)}"
            )
            raise

    async def _check_file_exists(self, bucket: str, path: str) -> bool:
        """Check if a file exists in storage without downloading it."""
        try:
            logger.debug(f"Checking if file exists in storage: {bucket}/{path}")
            
            # Use the storage service's check_file_exists method which is more reliable
            return await self.storage_service.check_file_exists(path, bucket)
        except Exception as e:
            logger.error(f"File existence check failed with error: {str(e)}", exc_info=True)
            return False  # Always assume file doesn't exist on error

# Service singleton
_document_service = None

def get_document_service() -> DocumentService:
    """Get the document service singleton."""
    global _document_service
    if _document_service is None:
        _document_service = DocumentService()
    return _document_service
